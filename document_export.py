"""
document_export.py
──────────────────
Assembles a python-docx Word document from a list of processed page results.

New API (mixed text + image regions):
    create_docx_document(page_data: list[dict]) -> io.BytesIO

    Each dict must have:
        "regions"      : list[RegionInfo]
        "original_bgr" : np.ndarray

Backward-compatible legacy API (plain-text pages):
    create_docx_document(page_texts: list[str]) -> io.BytesIO

Strict type-based insertion rules (Fix 5)
──────────────────────────────────────────
TEXT regions     → always typed paragraphs   — never image
EQUATION regions → always centred image      — never text   (caption: Equation N)
DIAGRAM regions  → always centred image      — never text   (caption: Figure N)

Never-empty guarantee:
  If a page produces zero usable regions after all validation, a fallback
  paragraph "(Page content could not be extracted)" is written instead of
  leaving that page completely blank in the document.
"""

from __future__ import annotations

import io
from typing import List, Union

import numpy as np
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from region_detector import RegionInfo
from structure_detector import TextBlock, detect_structure
from text_postprocessor import postprocess_text


# ─────────────────────────────────────────────────────────────────────────────
# Document-wide constants
# ─────────────────────────────────────────────────────────────────────────────

_PAGE_WIDTH_INCHES   = 6.0
_MAX_IMAGE_HEIGHT_IN = 7.0   # raised for large full-page diagrams
_CAPTION_FONT_SIZE   = Pt(9)
_NORMAL_SPACE_AFTER  = Pt(6)
_IMAGE_SPACE_BEFORE  = Pt(6)
_IMAGE_SPACE_AFTER   = Pt(6)
_CAPTION_SPACE_AFTER = Pt(10)


# ─────────────────────────────────────────────────────────────────────────────
# Style helpers
# ─────────────────────────────────────────────────────────────────────────────

def _apply_document_styles(document: Document) -> None:
    """Apply global font / colour settings and 1-inch page margins."""
    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    h1 = document.styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    h2 = document.styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    section = document.sections[0]
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(1.0)
    section.right_margin  = Inches(1.0)


# ─────────────────────────────────────────────────────────────────────────────
# TEXT rendering  (always typed text, never image)
# ─────────────────────────────────────────────────────────────────────────────

def _render_blocks(document: Document, blocks: List[TextBlock]) -> None:
    """Write TextBlocks as formatted paragraphs."""
    for block in blocks:
        if block.block_type == "heading1":
            document.add_heading(block.text, level=1)
        elif block.block_type == "heading2":
            document.add_heading(block.text, level=2)
        elif block.block_type == "bullet":
            para = document.add_paragraph(style="List Bullet")
            para.add_run(block.text)
        elif block.block_type == "numbered":
            para = document.add_paragraph(style="List Number")
            para.add_run(block.text)
        else:
            para = document.add_paragraph()
            para.add_run(block.text)
            para.paragraph_format.space_after = _NORMAL_SPACE_AFTER


def _insert_text_region(
    document: Document,
    region: RegionInfo,
    spell_check: bool = False,
) -> bool:
    """
    Insert a TEXT region as readable typed paragraphs.
    Returns True if anything was written, False if content was empty.
    Never inserts images.
    """
    raw = (region.content or "").strip()
    if not raw:
        return False

    processed = postprocess_text(raw, spell_check=spell_check)
    if not processed.strip():
        return False

    blocks = detect_structure(processed)
    if blocks:
        _render_blocks(document, blocks)
    else:
        para = document.add_paragraph()
        para.add_run(processed)
        para.paragraph_format.space_after = _NORMAL_SPACE_AFTER
    return True


# ─────────────────────────────────────────────────────────────────────────────
# IMAGE rendering  (equations & diagrams — always image, never text)
# ─────────────────────────────────────────────────────────────────────────────

def _insert_image_region(
    document: Document,
    region: RegionInfo,
    original_bgr: np.ndarray,
    counters: dict,
) -> bool:
    """
    Crop *region* from *original_bgr* and insert as a centred image + caption.
    Returns True on success, False if the crop was degenerate.
    Never writes region content as plain text.

    Fix 5: header text is always first (guaranteed by reading-order in
    region_detector); diagram image always follows below it.
    """
    from image_cropper import crop_region, region_to_png_bytes

    try:
        crop = crop_region(original_bgr, region)
        if crop is None or crop.size == 0:
            return False

        ch, cw = crop.shape[:2]
        if ch == 0 or cw == 0:
            return False

        png_bio = region_to_png_bytes(crop)

        # Compute display dimensions
        aspect    = cw / ch if ch else 1.0
        display_w = _PAGE_WIDTH_INCHES
        display_h = display_w / aspect

        if display_h > _MAX_IMAGE_HEIGHT_IN:
            display_h = _MAX_IMAGE_HEIGHT_IN
            display_w = min(display_h * aspect, _PAGE_WIDTH_INCHES)

        # Spacing before
        spacer = document.add_paragraph()
        spacer.paragraph_format.space_before = _IMAGE_SPACE_BEFORE
        spacer.paragraph_format.space_after  = Pt(2)

        # Centred image
        img_para = document.add_paragraph()
        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_para.paragraph_format.space_before = Pt(2)
        img_para.paragraph_format.space_after  = Pt(2)
        run = img_para.add_run()
        run.add_picture(png_bio, width=Inches(display_w))

        # Auto-numbered caption
        if region.type == "equation":
            counters["equation"] += 1
            caption_text = f"Equation {counters['equation']}"
        else:
            counters["figure"] += 1
            caption_text = f"Figure {counters['figure']}"

        cap_para = document.add_paragraph()
        cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_para.paragraph_format.space_before = Pt(2)
        cap_para.paragraph_format.space_after  = _CAPTION_SPACE_AFTER
        cap_run = cap_para.add_run(caption_text)
        cap_run.italic = True
        cap_run.font.size = _CAPTION_FONT_SIZE
        cap_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

        return True

    except Exception:
        # Placeholder label — not the equation math, never text content
        label = "[ Equation ]" if region.type == "equation" else "[ Figure / Diagram ]"
        para = document.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run  = para.add_run(label)
        run.italic = True
        run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
        return False


# ─────────────────────────────────────────────────────────────────────────────
# Page assembler — strict type dispatch with never-empty guarantee (Fix 5)
# ─────────────────────────────────────────────────────────────────────────────

def _assemble_region_page(
    document: Document,
    regions: List[RegionInfo],
    original_bgr: np.ndarray,
    counters: dict,
    spell_check: bool = False,
) -> None:
    """
    Assemble one page of regions into *document*.

    Dispatch (strict — no cross-type fallback):
        text     → _insert_text_region()     (always typed text)
        equation → _insert_image_region()    (always image)
        diagram  → _insert_image_region()    (always image)

    Fix 5 — Never-empty guarantee:
        Tracks whether anything was actually written.
        If not, inserts a placeholder so the page is never blank.
        For diagram-dominant pages the reading order from region_detector
        already guarantees header text appears before the diagram image.
    """
    wrote_anything = False

    for region in regions:
        if region.type == "text":
            ok = _insert_text_region(document, region, spell_check=spell_check)
        elif region.type in ("equation", "diagram"):
            ok = _insert_image_region(document, region, original_bgr, counters)
        else:
            ok = False
        if ok:
            wrote_anything = True

    # Never-empty guarantee
    if not wrote_anything:
        para = document.add_paragraph()
        run = para.add_run("(Page content could not be extracted)")
        run.italic = True
        run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)


def _assemble_legacy_page(
    document: Document,
    text: str,
    spell_check: bool = False,
) -> None:
    """Assemble a plain-text page (backward-compat path)."""
    page_text = text.strip()
    if not page_text:
        document.add_paragraph("(Empty page)")
        return
    processed = postprocess_text(page_text, spell_check=spell_check)
    blocks = detect_structure(processed)
    if blocks:
        _render_blocks(document, blocks)
    else:
        document.add_paragraph(processed)


# ─────────────────────────────────────────────────────────────────────────────
# Public API
# ─────────────────────────────────────────────────────────────────────────────

def create_docx_document(
    page_data: Union[List[dict], List[str]],
    spell_check: bool = False,
) -> io.BytesIO:
    """
    Build and return a .docx file as a BytesIO buffer.

    Parameters
    ----------
    page_data : list[dict] or list[str]
        New form  → each dict has 'regions' (list[RegionInfo])
                    and 'original_bgr' (np.ndarray).
        Legacy    → list of plain-text strings.
    spell_check : bool

    Returns
    -------
    io.BytesIO  — seeked to 0, ready for download.
    """
    document = Document()
    _apply_document_styles(document)

    counters = {"equation": 0, "figure": 0}

    if not page_data:
        document.add_paragraph("No content extracted.")
    else:
        is_new_api = isinstance(page_data[0], dict)

        for page_index, page in enumerate(page_data):
            if is_new_api:
                regions: List[RegionInfo] = page.get("regions", [])
                original_bgr: np.ndarray  = page.get(
                    "original_bgr", np.zeros((1, 1, 3), dtype=np.uint8)
                )
                _assemble_region_page(
                    document, regions, original_bgr,
                    counters, spell_check=spell_check,
                )
            else:
                _assemble_legacy_page(document, str(page), spell_check=spell_check)

            if page_index < len(page_data) - 1:
                document.add_page_break()

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer